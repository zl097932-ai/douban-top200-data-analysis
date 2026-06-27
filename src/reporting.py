"""中文 Word 报告、桌面 LibreOffice PDF 导出与页面渲染。"""

from __future__ import annotations

import inspect
import os
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


NAVY = "1F4D78"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
BODY_FONT = "宋体"
HEADING_FONT = BODY_FONT
CODE_FONT = BODY_FONT


def _set_run_font(
    run: Any,
    name: str = BODY_FONT,
    size: float | None = None,
    bold: bool | None = None,
    color: str | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def _set_style_font(style: Any, name: str, size: float, color: str, bold: bool, space_before: float, space_after: float, line_spacing: float) -> None:
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(space_before)
    style.paragraph_format.space_after = Pt(space_after)
    style.paragraph_format.line_spacing = line_spacing


def _add_page_number(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    _set_run_font(run, HEADING_FONT, 9, color="666666")
    field_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run._r.append(begin)
    field_run._r.append(instruction)
    field_run._r.append(end)
    run = paragraph.add_run(" 页")
    _set_run_font(run, HEADING_FONT, 9, color="666666")


def _set_cell_margins(cell: Any, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_widths(table: Any, widths_inches: list[float]) -> None:
    table.autofit = False
    table.allow_autofit = False
    widths_dxa = [round(width * 1440) for width in widths_inches]
    tbl_pr = table._tbl.tblPr
    table_width = tbl_pr.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        tbl_pr.append(table_width)
    table_width.set(qn("w:w"), str(sum(widths_dxa)))
    table_width.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    for grid_column, width_dxa in zip(table._tbl.tblGrid.gridCol_lst, widths_dxa):
        grid_column.set(qn("w:w"), str(width_dxa))
    for row in table.rows:
        for cell, width_dxa in zip(row.cells, widths_dxa):
            cell.width = Inches(width_dxa / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.first_child_found_in("w:tcW")
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:w"), str(width_dxa))
            tc_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _style_table(table: Any, widths_inches: list[float]) -> None:
    _set_table_widths(table, widths_inches)
    for cell in table.rows[0].cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), LIGHT_BLUE)
        cell._tc.get_or_add_tcPr().append(shading)
        for run in cell.paragraphs[0].runs:
            _set_run_font(run, HEADING_FONT, 13, bold=True, color=NAVY)
    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    _set_run_font(run, BODY_FONT, 12.5)
                paragraph.paragraph_format.line_spacing = 1.15


def _add_text(document: Document, text: str, *, align: int | None = None, first_line_indent: bool = True, italic: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.33
    paragraph.paragraph_format.space_after = Pt(8)
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    _set_run_font(run, BODY_FONT, 12)
    run.italic = italic


def _add_caption(document: Document, number: int, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(f"图 {number}  {text}")
    _set_run_font(run, BODY_FONT, 12, italic=False, color="555555")
    source = paragraph.add_run("\n数据来源：本研究基于豆瓣电影 Top250 公开列表页中排名前 200 名电影整理。")
    _set_run_font(source, BODY_FONT, 10.5, italic=False, color="666666")


def _add_figure(document: Document, path: Path, number: int, caption: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_picture(str(path), width=Cm(15.4))
    _add_caption(document, number, caption)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = True

    _set_style_font(document.styles["Normal"], BODY_FONT, 12, "000000", False, 0, 8, 1.33)
    _set_style_font(document.styles["Heading 1"], HEADING_FONT, 16, BLUE, True, 18, 10, 1.15)
    _set_style_font(document.styles["Heading 2"], HEADING_FONT, 13, BLUE, True, 12, 6, 1.15)
    _set_style_font(document.styles["Heading 3"], HEADING_FONT, 11, NAVY, True, 8, 4, 1.15)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("豆瓣电影 Top200 数据分析期末作业")
    _set_run_font(header_run, HEADING_FONT, 9, color="666666")
    _add_page_number(section.footer.paragraphs[0])


def _add_cover(document: Document, student: dict[str, str]) -> None:
    for _ in range(7):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Python 数据分析课程期末作业")
    _set_run_font(run, HEADING_FONT, 13, bold=True, color=BLUE)
    subtitle.paragraph_format.space_after = Pt(20)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("豆瓣电影 Top200\n数据爬取、清洗、分析与可视化")
    _set_run_font(run, HEADING_FONT, 24, bold=True, color=NAVY)

    description = document.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    description.paragraph_format.space_after = Pt(42)
    run = description.add_run("基于 Python 的公开电影榜单数据探索性分析报告")
    _set_run_font(run, BODY_FONT, 12, color="555555")

    metadata = [
        ("课程名称", student.get("course", "大数据技术基础")),
        ("学生姓名", student.get("name", "________")),
        ("学号", student.get("student_id", "________")),
        ("班级", student.get("class_name", "________")),
        ("完成日期", student.get("date", "2026 年 6 月")),
    ]
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(9)
        left = paragraph.add_run(f"{label}：")
        _set_run_font(left, HEADING_FONT, 12, bold=True)
        right = paragraph.add_run(value)
        _set_run_font(right, BODY_FONT, 12)
    document.add_page_break()


def _add_contents(document: Document) -> None:
    document.add_heading("目录", level=1)
    entries = [
        "摘要与关键词",
        "1 研究背景与目标",
        "2 数据来源与爬取方法",
        "3 数据清洗与质量核验",
        "4 探索性分析与可视化",
        "5 结论",
        "6 局限性",
        "参考文献",
        "附录 A 核心 Python 代码",
    ]
    for entry in entries:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.75)
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(entry)
        _set_run_font(run, BODY_FONT, 12)
    document.add_page_break()


def _quality_rows(quality: dict[str, Any]) -> list[tuple[str, str]]:
    return [
        ("原始记录数", str(quality["input_records"])),
        ("无效记录剔除数", str(quality["dropped_invalid_records"])),
        ("重复排名剔除数", str(quality["duplicate_rank_removed"])),
        ("重复片名/年份剔除数", str(quality["duplicate_title_year_removed"])),
        ("清洗后主表记录数", str(quality["output_records"])),
        ("清洗后缺失上映年份数", str(quality["missing_year_after_cleaning"])),
        ("评分范围", f"{quality['rating_range'][0]:.1f} - {quality['rating_range'][1]:.1f}"),
        ("排名是否唯一", "是" if quality["rank_is_unique"] else "否"),
    ]


def _add_quality_table(document: Document, quality: dict[str, Any]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "核验指标"
    table.rows[0].cells[1].text = "结果"
    for label, value in _quality_rows(quality):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    _style_table(table, [3.8, 2.7])


def _display_sample_value(value: Any) -> str:
    text = str(value).strip()
    return "—" if not text or text.lower() == "nan" else text


def _add_raw_sample_table(document: Document, records: list[dict[str, Any]]) -> None:
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["排名", "中文片名", "原始片名", "年份", "国家/地区", "评分"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for record in records:
        cells = table.add_row().cells
        values = [
            record.get("rank"),
            record.get("title"),
            record.get("original_title"),
            record.get("year"),
            record.get("country_text"),
            record.get("rating"),
        ]
        for cell, value in zip(cells, values):
            cell.text = _display_sample_value(value)
    _style_table(table, [0.7, 1.2, 1.65, 0.6, 1.65, 0.8])


def _add_code_block(document: Document, title: str, source: str) -> None:
    heading = document.add_heading(title, level=2)
    heading.paragraph_format.space_before = Pt(4)
    heading.paragraph_format.space_after = Pt(2)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 0.74
    run = paragraph.add_run(source)
    _set_run_font(run, CODE_FONT, 9.5)


def _format_p_value(value: float) -> str:
    return "< 0.001" if value < 0.001 else f"{value:.3f}"


def generate_report(
    output_docx: Path,
    summary: dict[str, Any],
    figures: dict[str, Path],
    student: dict[str, str],
    code_root: Path,
    raw_preview: list[dict[str, Any]],
) -> Path:
    """从真实统计摘要与图表生成中文课程报告。"""
    document = Document()
    _configure_document(document)
    _add_cover(document, student)
    _add_contents(document)

    overall = summary["overall"]
    quality = summary["quality"]
    top_genre = summary["top_genres"][0]
    top_country = summary["top_countries"][0]
    busiest_decade = max(summary["decades"], key=lambda item: item["movie_count"])
    best_decade = max(summary["decades"], key=lambda item: item["average_rating"])
    best_decade_count = int(best_decade["movie_count"])
    correlation = overall["rating_log_count_correlation"]
    pearson_p = _format_p_value(overall["rating_log_count_pearson_p_value"])
    rating_votes_rho = overall["rating_log_count_spearman_rho"]
    rating_votes_p = _format_p_value(overall["rating_log_count_spearman_p_value"])
    rank_rating_rho = overall["rank_rating_spearman_rho"]
    rank_rating_p = _format_p_value(overall["rank_rating_spearman_p_value"])
    report_date = str(summary.get("generated_at", ""))[:10] or "本次运行日期"

    document.add_heading("摘要", level=1)
    _add_text(
        document,
        f"本研究以豆瓣电影 Top250 公开列表页中的前 200 名为数据来源，使用 Python 完成分页数据采集、数据清洗、探索性分析和静态可视化。"
        f"本次运行共获得 {quality['input_records']} 条原始记录，清洗后保留 {overall['movie_count']} 条有效电影记录。"
        f"样本平均评分为 {overall['average_rating']:.2f} 分，中位数为 {overall['median_rating']:.2f} 分，覆盖 {overall['year_min']} 至 {overall['year_max']} 年上映的作品。"
        f"研究从类型、国家/地区、年代和用户评价规模等角度解释榜单构成，为理解高口碑电影的共同特征提供数据支持。",
    )
    keyword = document.add_paragraph()
    keyword.paragraph_format.space_after = Pt(10)
    label = keyword.add_run("关键词：")
    _set_run_font(label, HEADING_FONT, 12, bold=True)
    value = keyword.add_run("Python；网络爬虫；数据清洗；探索性分析；数据可视化；电影数据")
    _set_run_font(value, BODY_FONT, 12)

    document.add_heading("1 研究背景与目标", level=1)
    _add_text(document, "在线电影榜单同时包含作品属性、公众评分与评价人数等信息，是练习网络数据采集和探索性分析的合适场景。本作业以公开榜单页面为对象，避免登录、验证码绕过和高频访问，仅获取列表页中已公开展示的信息。")
    _add_text(document, "研究目标包括：第一，构建可复跑的 Python 数据管道；第二，通过规范化清洗使多字段文本能够用于统计；第三，比较榜单中不同类型、国家/地区和年代作品的分布；第四，观察评分、评价人数与榜单排名之间的关系。")

    document.add_heading("2 数据来源与爬取方法", level=1)
    _add_text(document, f"数据来自豆瓣电影 Top250 公开列表页（https://movie.douban.com/top250），本研究取其中排名第 1 至 200 名的电影作为 Top200 样本。本次数据处理日期为 {report_date}。程序依次访问 8 个分页，每页解析 25 条记录，采集排名、片名、原始片名、导演演员原始信息、上映年份、国家/地区、类型、评分、评价人数、短评和详情页链接。")
    _add_text(document, "该样本是榜单前列作品构成的非随机高口碑样本，适合描述榜单内部特征，不用于代表全部电影市场。为控制访问压力，程序在页与页之间随机等待 1.2 至 2.0 秒，对单页请求最多重试 3 次，并把每次请求结果写入 JSONL 日志。原始 HTML 页面和原始 CSV 会被保留；使用 --skip-fetch 时，程序不访问网络，而是基于已保存页面重新解析前 200 名数据，并生成明确标注为“离线复核”的 Top200 日志。全流程仅用于课程学习，未登录、未绕过验证码或访问限制。")

    document.add_heading("3 数据清洗与质量核验", level=1)
    _add_text(document, "清洗阶段首先将空白与缺失值统一为空字符串，避免将缺失原始片名错误保存为字符串“nan”；随后把排名、年份、评分和评价人数转换为数值类型，并验证评分范围、评价人数非负、排名唯一和年份合理性。导演和演员信息字段（people_info）保留列表页摘要原文，部分记录会以省略号截断；该字段仅用于原始信息追溯，不参与本报告的统计分析或可视化。对于国家/地区与类型等多值字段，主表保留原始组合文本，分析时另外拆分为一对多表，避免改变电影主表的记录粒度。")
    document.add_heading("3.1 原始数据样例", level=2)
    _add_text(document, "表 1 展示本次输入数据的前 5 条记录及主要字段；其中“—”表示网页中未提供原始片名，而不是字符串 nan。", first_line_indent=False)
    _add_raw_sample_table(document, raw_preview)
    sample_note = document.add_paragraph()
    sample_note.paragraph_format.space_before = Pt(4)
    sample_note.paragraph_format.space_after = Pt(8)
    sample_run = sample_note.add_run("表 1  原始数据样例（前 5 条）")
    _set_run_font(sample_run, BODY_FONT, 12, color="555555")
    document.add_heading("3.2 数据质量核验", level=2)
    _add_quality_table(document, quality)
    source_note = document.add_paragraph()
    source_note.paragraph_format.space_before = Pt(4)
    source_note.paragraph_format.space_after = Pt(8)
    run = source_note.add_run("表 2  数据清洗与质量核验结果（由本次程序运行自动生成）")
    _set_run_font(run, BODY_FONT, 12, color="555555")

    document.add_heading("4 探索性分析与可视化", level=1)
    figure_sections = [
        ("4.1 评分分布", "rating_distribution", "Top200 电影评分分布", f"样本评分集中在较高区间，平均值为 {overall['average_rating']:.2f} 分，说明该榜单筛选出的作品整体具有较高口碑。"),
        ("4.2 热门电影类型", "genre_counts", "Top200 中出现频次最高的电影类型", f"出现最多的类型是“{top_genre['genre']}”，共出现 {top_genre['count']} 次。由于一部电影可以具有多种类型，图中的总次数大于电影总数。"),
        ("4.3 主要制片国家/地区", "country_counts", "Top200 的主要制片国家/地区", f"“{top_country['country']}”出现 {top_country['count']} 次，是样本中出现频率最高的制片国家/地区。合拍片会分别计入其对应国家/地区。"),
        ("4.4 年代分布", "decade_counts", "不同年代入选 Top200 的电影数量", f"入选作品最多的是 {int(busiest_decade['decade'])} 年代，共 {int(busiest_decade['movie_count'])} 部，反映了该年代作品在榜单中的集中程度。"),
        ("4.5 年代与平均评分", "decade_ratings", "不同年代入选电影的平均评分", f"在样本覆盖的年代中，{int(best_decade['decade'])} 年代的平均评分最高，约为 {best_decade['average_rating']:.2f} 分；但该年代仅有 {best_decade_count} 部入选样本，样本量较小，均值仅供参考。该结果仅描述入选榜单的电影，不代表该年代所有电影的总体质量。"),
        ("4.6 评分与评价人数", "rating_votes", "评分与评价人数的关系", f"评分与评价人数对数的 Pearson 相关系数为 {correlation:.3f}（p={pearson_p}），Spearman 秩相关系数为 {rating_votes_rho:.3f}（p={rating_votes_p}）。两种结果均表明二者在本样本中仅呈极弱关系；图中的趋势线仅用于辅助观察，不代表因果关系。横轴取对数，是为了同时展示评价人数差异较大的作品。"),
        ("4.7 排名与评分", "rank_rating", "榜单排名与豆瓣评分的关系", f"排名与评分的 Spearman 秩相关系数为 {rank_rating_rho:.3f}（p={rank_rating_p}）。由于排名数字越小表示位置越靠前，负相关意味着较高评分通常对应更靠前的排名；趋势线用于展示总体方向，具体排名还会受到评分精度和评价人数等因素影响。"),
    ]
    for index, (heading, key, caption, explanation) in enumerate(figure_sections, start=1):
        document.add_heading(heading, level=2)
        _add_figure(document, figures[key], index, caption)
        _add_text(document, explanation)

    document.add_heading("5 结论", level=1)
    _add_text(document, f"（1）本项目成功构建了从网页采集到报告输出的 Python 自动化流程。清洗后 {overall['movie_count']} 条记录的排名均唯一，评分范围为 {overall['min_rating']:.1f} 至 {overall['max_rating']:.1f} 分，数据质量满足本次探索性分析需要。")
    _add_text(document, f"（2）Top200 呈现出明显的多类型特征，其中“{top_genre['genre']}”最常见；国家/地区方面，“{top_country['country']}”出现频次最高。两项分布均需结合榜单的口碑筛选性质理解，不能外推到全球全部电影市场。")
    _add_text(document, f"（3）时间维度上，{int(busiest_decade['decade'])} 年代入选作品最多，而 {int(best_decade['decade'])} 年代的样本平均评分最高。评分与评价人数对数的 Pearson 相关系数仅为 {correlation:.3f}（p={pearson_p}），说明在该 Top200 样本中没有明显的线性关联；排名与评分的秩相关更强，但这仅描述榜单内部的排序关系，不能据此推断因果。上述显著性检验用于辅助判断本样本内关系的稳定性，不应被解释为对全部电影总体的推断。")

    document.add_heading("6 局限性", level=1)
    _add_text(document, "第一，数据来自一个动态更新的公开榜单，抓取日期不同可能导致评分和评价人数变化。第二，样本为豆瓣 Top250 榜单的前 200 名，是非随机的高口碑选择样本，不能代表全部电影或电影市场。第三，国家/地区与类型来自页面文本拆分，合拍片和多类型电影会在展开统计中重复计入。第四，本项目的相关与显著性检验只用于描述和辅助解释本样本内的关系，不将相关关系解释为因果关系，也不将其外推为总体结论。")

    document.add_heading("参考文献", level=1)
    references = [
        f"[1] 豆瓣电影. 豆瓣电影 Top250（本研究取排名前 200 名）[EB/OL]. https://movie.douban.com/top250（访问日期：{report_date}）.",
        "[2] 张良均, 刘名奇, 龚志军, 等. Python 数据分析与挖掘实战（第2版）[M]. 北京: 机械工业出版社, 2019.",
        "[3] 李航. 统计学习方法（第2版）[M]. 北京: 清华大学出版社, 2019.",
        "[4] McKinney W. Python for Data Analysis[M]. 3rd ed. O'Reilly Media, 2022.",
        "[5] Hunter J D. Matplotlib: A 2D Graphics Environment[J]. Computing in Science & Engineering, 2007, 9(3): 90-95.",
    ]
    for reference in references:
        _add_text(document, reference, first_line_indent=False)

    document.add_heading("附录 A 核心 Python 代码", level=1)
    _add_text(document, "以下代码节选来自本项目源码，用于展示页面解析、低频分页采集、数据清洗、统计分析和可视化逻辑。完整可运行代码见项目的 src 目录；附录代码字号为 9.5 pt。", first_line_indent=False)
    from src.crawler import crawl_top200, parse_top250_html
    from src.processing import build_summary, clean_movies, create_figures

    _add_code_block(document, "A.1 单页 HTML 解析函数", inspect.getsource(parse_top250_html))
    _add_code_block(document, "A.2 分页采集函数", inspect.getsource(crawl_top200))
    _add_code_block(document, "A.3 清洗函数", inspect.getsource(clean_movies))
    _add_code_block(document, "A.4 统计分析函数", inspect.getsource(build_summary))
    _add_code_block(document, "A.5 可视化函数", inspect.getsource(create_figures))

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx)
    return output_docx


def find_desktop_soffice(explicit_path: str | None = None) -> Path:
    """优先定位用户指定的桌面版 LibreOffice，再查找常见安装路径。"""
    candidates = [
        explicit_path,
        os.environ.get("LIBREOFFICE_SOFFICE"),
        str(Path.home() / "Tools" / "LibreOfficePortable" / "App" / "libreoffice" / "program" / "soffice.com"),
        str(Path(os.environ.get("ProgramFiles", r"C:\\Program Files")) / "LibreOffice" / "program" / "soffice.com"),
        shutil.which("soffice.com"),
        shutil.which("soffice"),
    ]
    for item in candidates:
        if item and Path(item).is_file():
            return Path(item)
    raise FileNotFoundError("未找到桌面版 LibreOffice。请安装 LibreOffice，或使用 --soffice 指定 soffice.com 路径。")


def export_pdf_with_desktop_libreoffice(docx_path: Path, pdf_dir: Path, explicit_soffice: str | None = None) -> Path:
    """使用桌面版 LibreOffice 无界面导出 PDF，避免替代渲染器的兼容差异。"""
    soffice = find_desktop_soffice(explicit_soffice)
    pdf_dir.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="douban_lo_profile_") as profile_dir:
        profile_uri = Path(profile_dir).resolve().as_uri()
        command = [
            str(soffice),
            "--headless",
            f"-env:UserInstallation={profile_uri}",
            "--convert-to",
            "pdf:writer_pdf_Export",
            "--outdir",
            str(pdf_dir.resolve()),
            str(docx_path.resolve()),
        ]
        completed = subprocess.run(command, capture_output=True, text=True, encoding="utf-8", errors="replace", timeout=180, check=False)
    pdf_path = pdf_dir / f"{docx_path.stem}.pdf"
    if completed.returncode != 0 or not pdf_path.exists() or pdf_path.stat().st_size == 0:
        raise RuntimeError(
            "桌面 LibreOffice PDF 导出失败。"
            f"\n命令输出：{completed.stdout}\n错误输出：{completed.stderr}"
        )
    return pdf_path


def render_pdf_pages(pdf_path: Path, output_dir: Path, dpi: int = 150) -> list[Path]:
    """将桌面 LibreOffice 导出的 PDF 转为 PNG，供逐页视觉检查。"""
    output_dir.mkdir(parents=True, exist_ok=True)
    for old_image in output_dir.glob("page-*.png"):
        old_image.unlink()
    document = fitz.open(pdf_path)
    scale = dpi / 72
    image_paths: list[Path] = []
    for index, page in enumerate(document, start=1):
        pixmap = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
        output_path = output_dir / f"page-{index:02d}.png"
        pixmap.save(str(output_path))
        image_paths.append(output_path)
    document.close()
    if not image_paths:
        raise RuntimeError("PDF 没有可渲染页面。")
    return image_paths
