"""Build the GitHub showcase DOCX report for the Douban Top200 project."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PROCESSED = ROOT / "data" / "processed"
FIGURES = ROOT / "output" / "figures"
REPORT = ROOT / "output" / "报告" / "豆瓣电影Top200数据分析报告_GitHub展示版.docx"

BLACK = "000000"
LIGHT_GRAY = "F2F2F2"
BODY = BLACK
MUTED = BLACK
FONT_CJK = "宋体"
MODEL_DISPLAY_NAMES = {
    "Baseline mean rating": "平均分基线",
    "Ridge regression": "岭回归",
    "Random forest": "随机森林",
}


def set_rfonts(r_pr) -> None:
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key in list(r_fonts.attrib):
        local_name = key.rsplit("}", 1)[-1].lower()
        if local_name.endswith("theme"):
            del r_fonts.attrib[key]
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT_CJK)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = FONT_CJK
    set_rfonts(run._element.get_or_add_rPr())
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color or BLACK)


def set_style(style, size: float, color: str, bold: bool, before: float, after: float, line_spacing: float) -> None:
    style.font.name = FONT_CJK
    r_pr = style._element.rPr
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        style._element.append(r_pr)
    set_rfonts(r_pr)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line_spacing


def set_cell_fill(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]) -> None:
    table.autofit = False
    table.allow_autofit = False
    widths_dxa = [round(width * 1440) for width in widths]
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")

    for grid_column, width_dxa in zip(table._tbl.tblGrid.gridCol_lst, widths_dxa):
        grid_column.set(qn("w:w"), str(width_dxa))

    for row in table.rows:
        for cell, width_dxa in zip(row.cells, widths_dxa):
            cell.width = Inches(width_dxa / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_dxa))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def style_table(table, widths: list[float]) -> None:
    set_table_geometry(table, widths)
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if i == 0:
                set_cell_fill(cell, LIGHT_GRAY)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    set_run_font(run, 10.5, bold=(i == 0), color=BLACK)


def add_paragraph(doc: Document, text: str, *, style: str | None = None, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, 11, True, BLACK)
        rest = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(rest, 11, False, BODY)
        return
    run = paragraph.add_run(text)
    set_run_font(run, 11, False, BODY)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.1
        run = paragraph.add_run(item)
        set_run_font(run, 10.5, False, BODY)


def add_report_heading(doc: Document, text: str, level: int = 1) -> None:
    sizes = {1: 16, 2: 13, 3: 12}
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.space_before = Pt({1: 16, 2: 12, 3: 8}.get(level, 8))
    paragraph.paragraph_format.space_after = Pt({1: 8, 2: 6, 3: 4}.get(level, 4))
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    set_run_font(run, sizes.get(level, 12), True, BLACK)


def add_table_of_contents(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("目录")
    set_run_font(run, 16, True, BLACK)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "请在 Word 中右键更新目录，或使用 Ctrl+A 后按 F9 更新全部域。"

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(separate)
    run._r.append(placeholder)
    run._r.append(end)
    set_run_font(run, 11, False, BLACK)


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.insert(0, update_fields)
    update_fields.set(qn("w:val"), "true")


def add_key_value_table(doc: Document, rows: list[tuple[str, str]], widths: list[float] | None = None) -> None:
    widths = widths or [1.8, 4.7]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "说明"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    style_table(table, widths)


def add_matrix_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
    style_table(table, widths)


def format_int(value: float | int) -> str:
    return f"{int(round(value)):,}"


def format_pct(part: float | int, whole: float | int) -> str:
    return f"{part / whole * 100:.1f}%"


def format_p_value(value: float) -> str:
    return "<0.001" if value < 0.001 else f"{value:.3f}"


def add_category_matrix(doc: Document, genres: pd.Series, countries: pd.Series) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["类型", "出现次数", "国家/地区", "出现次数"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    pairs = list(zip(genres.items(), countries.items()))
    for (genre, genre_count), (country, country_count) in pairs:
        cells = table.add_row().cells
        cells[0].text = str(genre)
        cells[1].text = f"{genre_count} 次"
        cells[2].text = str(country)
        cells[3].text = f"{country_count} 次"
    style_table(table, [1.8, 1.2, 2.1, 1.4])


def add_ml_model_table(doc: Document, models: dict[str, dict[str, float]]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["模型", "MAE", "RMSE", "R²"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for model_name, values in models.items():
        cells = table.add_row().cells
        cells[0].text = MODEL_DISPLAY_NAMES.get(model_name, model_name)
        cells[1].text = f"{values['mae_mean']:.4f}"
        cells[2].text = f"{values['rmse_mean']:.4f}"
        cells[3].text = f"{values['r2_mean']:.4f}"
    style_table(table, [2.4, 1.3, 1.3, 1.5])


def add_cluster_profile_table(doc: Document, profiles: list[dict]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["分群画像", "数量", "平均评分", "代表影片"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for profile in profiles:
        cells = table.add_row().cells
        cells[0].text = str(profile["profile"])
        cells[1].text = str(profile["movie_count"])
        cells[2].text = f"{profile['average_rating']:.3f}"
        cells[3].text = str(profile["example_titles"])
    style_table(table, [1.6, 0.8, 0.9, 3.2])


def add_feature_importance_table(doc: Document, feature_importance: list[dict]) -> None:
    rows = [
        [
            str(item["display_name"]),
            f"{item['importance']:.4f}",
            f"{item['importance_std']:.4f}",
        ]
        for item in feature_importance[:8]
    ]
    add_matrix_table(doc, ["特征", "重要性均值", "标准差"], rows, [3.5, 1.5, 1.5])


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, 9, False, MUTED)
    fld_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    fld_run._r.append(begin)
    fld_run._r.append(instruction)
    fld_run._r.append(end)
    run = paragraph.add_run(" 页")
    set_run_font(run, 9, False, MUTED)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    set_run_font(run, 9, False, MUTED)


def load_data() -> dict:
    movies = pd.read_csv(PROCESSED / "movies_clean.csv", encoding="utf-8-sig")
    genres = pd.read_csv(PROCESSED / "movies_by_genre.csv", encoding="utf-8-sig")
    countries = pd.read_csv(PROCESSED / "movies_by_country.csv", encoding="utf-8-sig")
    decades = pd.read_csv(PROCESSED / "decade_summary.csv", encoding="utf-8-sig")
    analysis_summary = json.loads((PROCESSED / "analysis_summary.json").read_text(encoding="utf-8"))
    ml_summary_path = PROCESSED / "ml_summary.json"
    ml_summary = json.loads(ml_summary_path.read_text(encoding="utf-8")) if ml_summary_path.exists() else None
    rating_labels = ["8.5-8.7", "8.8-9.0", "9.1-9.3", "9.4-9.6", "9.7 及以上"]
    rating_bins = (
        pd.cut(
            movies["rating"],
            bins=[8.49, 8.7, 9.0, 9.3, 9.6, 9.8],
            labels=rating_labels,
            include_lowest=True,
        )
        .value_counts()
        .reindex(rating_labels)
    )
    genre_profile = (
        genres.groupby("genre")
        .agg(movie_count=("title", "count"), average_rating=("rating", "mean"))
        .query("movie_count >= 10")
        .sort_values(["movie_count", "average_rating"], ascending=[False, False])
    )
    country_profile = (
        countries.groupby("country")
        .agg(movie_count=("title", "count"), average_rating=("rating", "mean"))
        .query("movie_count >= 5")
        .sort_values(["movie_count", "average_rating"], ascending=[False, False])
    )
    return {
        "movies": movies,
        "genres": genres,
        "countries": countries,
        "decades": decades,
        "analysis_summary": analysis_summary,
        "top_genres": genres["genre"].value_counts().head(8),
        "top_countries": countries["country"].value_counts().head(8),
        "rating_bins": rating_bins,
        "genre_profile": genre_profile,
        "country_profile": country_profile,
        "popular_movies": movies.sort_values("rating_count", ascending=False).head(8),
        "ml_summary": ml_summary,
    }


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    set_style(doc.styles["Normal"], 11, BLACK, False, 0, 6, 1.1)
    set_style(doc.styles["Heading 1"], 16, BLACK, True, 16, 8, 1.1)
    set_style(doc.styles["Heading 2"], 13, BLACK, True, 12, 6, 1.1)
    set_style(doc.styles["Heading 3"], 12, BLACK, True, 8, 4, 1.1)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def build_report() -> Path:
    data = load_data()
    movies = data["movies"]
    decades = data["decades"]
    ml_summary = data["ml_summary"]
    summary = data["analysis_summary"]
    quality = summary["quality"]
    overall = summary["overall"]
    movie_count = len(movies)
    rating_stats = movies["rating"].describe()
    rating_count_stats = movies["rating_count"].describe()
    top_decade = decades.sort_values("movie_count", ascending=False).iloc[0]

    doc = Document()
    configure_document(doc)
    doc.core_properties.author = "刘智鹏"
    doc.core_properties.title = "豆瓣电影 Top200 数据分析报告"
    doc.core_properties.subject = "Python 数据采集、清洗、分析与可视化"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("豆瓣电影 Top200 数据分析报告")
    set_run_font(run, 24, True, BLACK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("基于公开榜单数据的描述统计、分组比较与可复现建模")
    set_run_font(run, 12, False, BLACK)

    add_table_of_contents(doc)
    doc.add_section(WD_SECTION_START.NEW_PAGE)

    add_report_heading(doc, "1. 研究对象与结论摘要", level=1)
    add_paragraph(
        doc,
        "本报告使用仓库内保存的豆瓣电影 Top250 公开榜单前 200 部影片作为样本，分析对象为影片排名、上映年份、国家/地区、类型、评分和评价人数等字段。数据口径为一部电影一条主记录；类型和国家/地区属于多值字段，在分布统计中按展开后的出现次数计算。",
    )
    add_paragraph(
        doc,
        f"总体来看，样本评分均值为 {overall['average_rating']:.2f}，中位数为 {overall['median_rating']:.2f}，四分位区间为 {rating_stats['25%']:.1f}-{rating_stats['75%']:.1f}。这说明 Top200 榜单是一个明显的高分截断样本，分析重点不应放在“是否高分”，而应放在高分影片内部的结构差异。",
    )
    add_key_value_table(
        doc,
        [
            ("电影数量", f"{movie_count} 部"),
            ("年份跨度", f"{int(movies['year'].min())}-{int(movies['year'].max())}"),
            ("平均评分 / 中位评分", f"{overall['average_rating']:.2f} / {overall['median_rating']:.2f}"),
            ("评分标准差", f"{rating_stats['std']:.2f}"),
            ("评分范围", f"{overall['min_rating']:.1f}-{overall['max_rating']:.1f}"),
            ("累计评价人数", format_int(overall["total_rating_count"])),
            ("评价人数中位数", format_int(rating_count_stats["50%"])),
            ("出现最多的类型", str(data["top_genres"].index[0])),
            ("出现最多的国家/地区", str(data["top_countries"].index[0])),
            ("数量最多的年代", f"{int(top_decade['decade'])} 年代（{int(top_decade['movie_count'])} 部）"),
        ],
    )

    add_report_heading(doc, "2. 数据质量与统计口径", level=1)
    add_paragraph(
        doc,
        "清洗过程保留了 200 条有效记录，排名字段唯一，年份、国家/地区和类型字段均无清洗后缺失。由于榜单电影常同时属于多个类型或多个制片地区，后续类型与地区统计采用展开表：同一部电影可能在多个标签下各计一次，因此这些次数用于描述结构，不等同于互斥分类占比。",
    )
    add_key_value_table(
        doc,
        [
            ("原始记录数", f"{quality['input_records']} 条"),
            ("清洗后记录数", f"{quality['output_records']} 条"),
            ("无效记录剔除数", f"{quality['dropped_invalid_records']} 条"),
            ("重复排名剔除数", f"{quality['duplicate_rank_removed']} 条"),
            ("重复片名年份剔除数", f"{quality['duplicate_title_year_removed']} 条"),
            ("年份缺失", f"{quality['missing_year_after_cleaning']} 条"),
            ("国家/地区缺失", f"{quality['missing_country_after_cleaning']} 条"),
            ("类型缺失", f"{quality['missing_genre_after_cleaning']} 条"),
            ("排名唯一性", "通过" if quality["rank_is_unique"] else "需检查"),
        ],
    )

    add_report_heading(doc, "3. 评分分布分析", level=1)
    add_paragraph(
        doc,
        f"评分分布集中在 8.8-9.3 分之间。样本最低分为 {overall['min_rating']:.1f}，最高分为 {overall['max_rating']:.1f}，标准差只有 {rating_stats['std']:.2f}，说明榜单内部评分差异相对有限。若只看均值，容易忽略这种“高分样本内部再排序”的特点，因此下面按评分区间拆分观察。",
    )
    rating_rows = [
        [str(label), f"{int(count)} 部", format_pct(int(count), movie_count)]
        for label, count in data["rating_bins"].items()
    ]
    add_matrix_table(doc, ["评分区间", "电影数量", "占样本比例"], rating_rows, [2.1, 1.6, 2.8])
    add_paragraph(
        doc,
        "从区间分布看，8.8-9.0 分影片最多，9.1-9.3 分次之，9.4 分以上影片数量明显减少。这种分布说明 Top200 影片整体质量接近，但头部高分影片仍具有较强区分度。",
    )

    add_report_heading(doc, "4. 类型、地区与年代分析", level=1)
    findings = [
        f"剧情片出现 {int(data['top_genres'].iloc[0])} 次，是样本中最突出的类型标签。",
        f"美国电影出现 {int(data['top_countries'].iloc[0])} 次，是国家/地区来源中的最高频标签。",
        f"{int(top_decade['decade'])} 年代影片数量最多，共 {int(top_decade['movie_count'])} 部；1990 年代和 2010 年代也占据较大比例。",
        "类型和地区平均分只能解释样本内部差异，不能直接推出某一类型或地区整体质量更高。",
    ]
    add_bullets(doc, findings)
    add_category_matrix(doc, data["top_genres"], data["top_countries"])

    genre_rows = [
        [genre, f"{int(row.movie_count)} 次", f"{row.average_rating:.3f}"]
        for genre, row in data["genre_profile"].head(8).iterrows()
    ]
    add_paragraph(doc, "按出现次数较多的类型继续比较平均评分，可以看到剧情、家庭、科幻等标签的平均分略高，但差距整体较小。由于电影类型是多标签字段，同一电影可能同时贡献给多个类型，本表更适合观察结构，而不是做因果判断。")
    add_matrix_table(doc, ["类型", "出现次数", "平均评分"], genre_rows, [2.3, 1.7, 2.5])

    country_rows = [
        [country, f"{int(row.movie_count)} 次", f"{row.average_rating:.3f}"]
        for country, row in data["country_profile"].head(8).iterrows()
    ]
    add_paragraph(doc, "国家/地区分布中，美国样本量最大，中国大陆、意大利、德国等标签的样本平均分较高，但其中部分地区样本数较小，解读时需要同时看数量和均值。")
    add_matrix_table(doc, ["国家/地区", "出现次数", "平均评分"], country_rows, [2.3, 1.7, 2.5])

    decade_rows = [
        [f"{int(row.decade)} 年代", f"{int(row.movie_count)} 部", f"{row.average_rating:.3f}"]
        for row in decades.sort_values("decade").itertuples()
    ]
    add_paragraph(doc, "年代分布显示，1990 年代、2000 年代和 2010 年代构成榜单主体。早期年代的平均分较高，但样本量明显偏少，因此更适合作为经典影片留存现象的提示，而不是年代优劣比较。")
    add_matrix_table(doc, ["年代", "电影数量", "平均评分"], decade_rows, [2.1, 1.6, 2.8])

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_report_heading(doc, "5. 排名、评价人数与评分关系", level=1)
    add_paragraph(
        doc,
        f"排名与评分的 Spearman 秩相关系数为 {overall['rank_rating_spearman_rho']:.3f}，p 值为 {format_p_value(overall['rank_rating_spearman_p_value'])}。由于排名数字越小表示位置越靠前，负相关说明评分越高的电影通常排名越靠前，这与榜单排序逻辑一致。",
    )
    add_paragraph(
        doc,
        f"评分与评价人数对数之间的相关性较弱：Pearson 相关系数为 {overall['rating_log_count_correlation']:.3f}，p 值为 {format_p_value(overall['rating_log_count_pearson_p_value'])}；Spearman 相关系数为 {overall['rating_log_count_spearman_rho']:.3f}，p 值为 {format_p_value(overall['rating_log_count_spearman_p_value'])}。因此，在该样本中，评价人数多并不必然对应更高评分。",
    )
    popular_rows = [
        [str(int(row.rank)), str(row.title), f"{row.rating:.1f}", format_int(row.rating_count)]
        for row in data["popular_movies"].itertuples()
    ]
    add_matrix_table(doc, ["排名", "影片", "评分", "评价人数"], popular_rows, [0.8, 3.1, 0.9, 1.7])

    if ml_summary:
        add_report_heading(doc, "6. 机器学习补充分析", level=1)
        model_summary = ml_summary["model_summary"]
        relative_improvement = model_summary["mae_improvement_vs_baseline"] / model_summary["baseline_mae"] * 100
        add_paragraph(
            doc,
            "机器学习部分只作为补充分析，用来检验现有结构化特征是否能解释一部分评分差异。特征包括年份、距今年数、评价人数对数、类型标签数、国家/地区数，以及热门类型和国家/地区的多热编码。模型采用 5 折交叉验证，并与平均分基线对照。",
        )
        add_ml_model_table(doc, model_summary["models"])
        add_paragraph(
            doc,
            f"随机森林的 MAE 为 {model_summary['best_mae']:.4f}，比平均分基线低 {model_summary['mae_improvement_vs_baseline']:.4f}，相对改善约 {relative_improvement:.1f}%。这说明特征中确实包含一定解释信息，但样本只有 200 条，且评分范围窄，模型结果应理解为探索性证据。",
        )
        add_paragraph(doc, "置换特征重要性显示，评价人数（对数）、上映年份、距今年数和类型标签数对误差影响较大。该结果与前面的描述统计相互补充：热度、年代和内容结构会影响榜单内部评分差异，但不能解释为因果关系。")
        add_feature_importance_table(doc, ml_summary["feature_importance"])
        add_paragraph(doc, "KMeans 分群将影片划分为若干内容画像，便于从结构上理解榜单组成。分群结果不是人工分类标准，而是根据年份、热度、类型和地区等特征形成的相似性分组。")
        add_cluster_profile_table(doc, ml_summary["cluster_profiles"])

        doc.add_section(WD_SECTION_START.NEW_PAGE)

    add_report_heading(doc, "7. 可视化结果", level=1)
    figures = [
        ("01_评分分布.png", "图 1 Top200 电影评分分布"),
        ("02_热门电影类型.png", "图 2 热门电影类型出现次数"),
        ("03_主要制片国家地区.png", "图 3 主要制片国家/地区出现次数"),
        ("04_不同年代电影数量.png", "图 4 不同年代电影数量"),
        ("07_排名与评分关系.png", "图 5 排名与评分关系"),
    ]
    if ml_summary:
        figures.extend(
            [
                ("08_机器学习评分预测对比.png", "图 6 机器学习评分预测 MAE 对比"),
                ("09_机器学习特征重要性.png", "图 7 随机森林置换特征重要性"),
                ("10_电影画像分群.png", "图 8 电影画像分群二维投影"),
            ]
        )
    for file_name, caption in figures:
        doc.add_picture(str(FIGURES / file_name), width=Inches(5.9))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, caption)

    add_report_heading(doc, "8. 项目输出与可复现操作", level=1)
    add_paragraph(doc, "仓库保留了原始数据、清洗数据、统计摘要、机器学习摘要、图表和报告生成脚本。复现时既可以使用仓库内保存的原始 CSV 离线生成结果，也可以重新访问公开榜单页面刷新数据。")
    add_key_value_table(
        doc,
        [
            ("清洗数据", "data/processed/movies_clean.csv"),
            ("类型展开表", "data/processed/movies_by_genre.csv"),
            ("国家/地区展开表", "data/processed/movies_by_country.csv"),
            ("统计摘要", "data/processed/analysis_summary.json"),
            ("机器学习摘要", "data/processed/ml_summary.json"),
            ("电影分群结果", "data/processed/movie_ml_clusters.csv"),
            ("分析图表", "output/figures/*.png"),
            ("展示报告", "GitHub 展示版 DOCX/PDF"),
        ],
    )

    add_key_value_table(
        doc,
        [
            ("安装依赖", "python -m venv .venv；.venv\\Scripts\\python.exe -m pip install -r requirements.txt"),
            ("离线复现", "python run_pipeline.py --skip-fetch --student-name ... --student-id ... --class-name ..."),
            ("重新采集", "python run_pipeline.py --fetch --student-name ... --student-id ... --class-name ..."),
            ("运行 ML 扩展", "python scripts/run_ml_extension.py"),
            ("重建展示报告", "python scripts/build_github_showcase_report.py"),
            ("运行测试", "python -m pytest -q"),
        ],
    )

    add_report_heading(doc, "9. 项目实现说明", level=1)
    add_key_value_table(
        doc,
        [
            ("采集层", "分页访问公开榜单页面，解析排名、片名、年份、国家/地区、类型、评分和评价人数。"),
            ("清洗层", "统一字段类型，处理多值字段，移除重复项，并输出一电影一行的主表。"),
            ("分析层", "生成评分、类型、国家/地区、年代和相关关系统计。"),
            ("建模层", "执行评分预测交叉验证、随机森林特征重要性分析和 KMeans 电影分群。"),
            ("输出层", "保存 CSV/JSON、PNG 图表、Word 报告和 PDF 报告。"),
        ],
    )

    add_report_heading(doc, "10. 局限性与改进方向", level=1)
    add_bullets(
        doc,
        [
            "样本来自豆瓣 Top250 前 200 名，是高口碑影片集合，不代表全部电影市场，也不能用于估计普通电影评分分布。",
            "榜单页面会随时间变化，本报告反映当前仓库数据快照。若重新采集，评分、评价人数和排名可能发生变化。",
            "类型与国家/地区为多值字段，展开统计会重复计入同一电影，因此不应把类型或地区次数相加后解释为互斥比例。",
            "相关系数和机器学习结果只说明样本内特征关系，不能证明评分由某个类型、地区或年代直接导致。",
            "后续可以补充评论文本情感分析、时间序列更新、交互式看板和 GitHub Actions 自动复现实验流程。",
        ],
    )

    REPORT.parent.mkdir(parents=True, exist_ok=True)
    set_update_fields_on_open(doc)
    doc.save(REPORT)
    return REPORT


if __name__ == "__main__":
    print(build_report())
